# -*- coding: utf-8 -*-
"""
Build a professional Word newsletter for ARK Capture Solutions.

Document structure per section:
  ├── Market Background  (baseline bullets from ark-baseline.yaml)
  ├── Updates This Month (GPT-4o article entries from digest markdown)
  └── Changes Since Last Issue (from digest markdown)

Designed to be sent directly to ARK's leadership team (CEO/CTO/BD).
Branding: "Prepared by GG Advisory for ARK Capture Solutions"

Usage:
    python -m src.build_ark_newsletter out/ark/monthly-digest-2026-03.md

Env vars:
    CFG_BASELINE=config/ark-baseline.yaml
"""
from __future__ import annotations

import os
import re
import sys
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

import yaml
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

CFG_BASELINE = os.getenv("CFG_BASELINE", "config/ark-baseline.yaml")

# ── Brand palette ─────────────────────────────────────────────────────────────
GG_TEAL      = RGBColor(0x1B, 0x7A, 0x6B)   # GG Advisory teal
GG_NAVY      = RGBColor(0x0D, 0x2D, 0x4A)   # deep navy for headings
DARK_GREY    = RGBColor(0x2C, 0x2C, 0x2C)   # body text
MID_GREY     = RGBColor(0x55, 0x55, 0x55)   # labels / dates
LIGHT_GREY   = RGBColor(0x88, 0x88, 0x88)   # footer text
LINK_BLUE    = RGBColor(0x00, 0x5E, 0xA2)   # hyperlinks
ACCENT_AMBER = RGBColor(0xE0, 0x8A, 0x00)   # "Why it matters" label
LIGHT_TEAL_BG = "D6EDE9"                     # shading hex for placeholder boxes

# Priority tag palette
_TAG_COLORS: Dict[str, Tuple[str, str]] = {
    "teal":   ("1B7A6B", "FFFFFF"),
    "navy":   ("0D2D4A", "FFFFFF"),
    "red":    ("C0392B", "FFFFFF"),
    "orange": ("E67E22", "FFFFFF"),
    "grey":   ("888888", "FFFFFF"),
}

_SECTION_DISPLAY: Dict[str, str] = {
    "grants_funding":  "Grants & Funding",
    "market_policy":   "Market & Policy",
    "competitors":     "Competitors",
    "partners_buyers": "Partners & Buyers",
}

_SECTION_ORDER = ["grants_funding", "market_policy", "competitors", "partners_buyers"]

# Reverse map: old-format display name → section key (backward compat with pre-v1.1 digests)
_LEGACY_SECTION_MAP: Dict[str, str] = {
    "grants & funding":  "grants_funding",
    "market & policy":   "market_policy",
    "competitors":       "competitors",
    "partners & buyers": "partners_buyers",
}


# ── Low-level helpers ─────────────────────────────────────────────────────────

def _cm_to_emu(cm: float) -> int:
    return int(cm * 914400 / 2.54)


def _font(run, size_pt: float, bold=False, italic=False,
          color: Optional[RGBColor] = None, underline=False):
    run.font.name = "Calibri"
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    run.font.underline = underline
    if color:
        run.font.color.rgb = color


def _para(doc, space_before=0, space_after=6, left_indent=0,
          alignment=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after  = Pt(space_after)
    pf.alignment    = alignment
    if left_indent:
        pf.left_indent = Inches(left_indent)
    return p


def _hrule(doc, color_hex: str = "1B7A6B", thickness: int = 12):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    str(thickness))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def _shaded_para(doc, text: str, bg_hex: str = LIGHT_TEAL_BG,
                 italic=True, size_pt=10):
    """A paragraph with a solid background colour (e.g. light teal box)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.left_indent  = Inches(0.15)
    p.paragraph_format.right_indent = Inches(0.15)
    # Apply shading to paragraph
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  bg_hex)
    pPr.append(shd)
    r = p.add_run(text)
    _font(r, size_pt, italic=italic, color=DARK_GREY)
    return p


def _hyperlink(paragraph, url: str, display: str):
    """Insert a clickable hyperlink run into an existing paragraph."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hl  = OxmlElement("w:hyperlink")
    hl.set(qn("r:id"),      r_id)
    hl.set(qn("w:history"), "1")
    run  = OxmlElement("w:r")
    rPr  = OxmlElement("w:rPr")
    color_el = OxmlElement("w:color")
    color_el.set(qn("w:val"), "005EA2")
    rPr.append(color_el)
    for sz_tag in ("w:sz", "w:szCs"):
        sz = OxmlElement(sz_tag)
        sz.set(qn("w:val"), "20")
        rPr.append(sz)
    run.append(rPr)
    t = OxmlElement("w:t")
    t.text = display
    run.append(t)
    hl.append(run)
    paragraph._p.append(hl)


def _tag_run(paragraph, tag_text: str, bg_hex: str, fg_hex: str):
    """Inline coloured tag chip (simulated via shaded run)."""
    r = paragraph.add_run(f"  {tag_text}  ")
    r.font.name  = "Calibri"
    r.font.size  = Pt(8)
    r.font.bold  = True
    rPr = r._r.get_or_add_rPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  bg_hex)
    rPr.append(shd)
    color_el = OxmlElement("w:color")
    color_el.set(qn("w:val"), fg_hex)
    rPr.append(color_el)


def _left_border_para(doc, space_before=10, space_after=4, border_color="1B7A6B"):
    """Paragraph with a teal left border for subsection headings."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.left_indent  = Inches(0.15)
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"),   "single")
    left.set(qn("w:sz"),    "18")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), border_color)
    pBdr.append(left)
    pPr.append(pBdr)
    return p


# ── Section & content builders ────────────────────────────────────────────────

def _section_heading(doc, display_name: str):
    """Top-level section heading (navy, bold) with a rule below."""
    _para(doc, space_before=16, space_after=0)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(display_name.upper())
    _font(r, 12, bold=True, color=GG_NAVY)
    _hrule(doc, color_hex="0D2D4A", thickness=10)


def _subsection_heading(doc, title: str, size_pt=10.5,
                        color: Optional[RGBColor] = None):
    """Subsection heading with teal left border."""
    p = _left_border_para(doc, space_before=10, space_after=4)
    r = p.add_run(title)
    _font(r, size_pt, bold=True, color=color or GG_TEAL)
    return p


def _baseline_entry(doc, entry: Dict):
    """Render one baseline entry: label (+ tag chip) + bullet list."""
    # Label line
    p = _para(doc, space_before=8, space_after=2)
    r = p.add_run(entry.get("label", ""))
    _font(r, 10.5, bold=True, color=DARK_GREY)

    tag = entry.get("tag")
    if tag:
        tag_color_key = entry.get("tag_color", "teal")
        bg_hex, fg_hex = _TAG_COLORS.get(tag_color_key, _TAG_COLORS["teal"])
        p.add_run("  ")
        _tag_run(p, tag, bg_hex, fg_hex)

    # Bullet points
    for bullet in entry.get("bullets", []):
        pb = doc.add_paragraph(style="List Bullet")
        pb.paragraph_format.space_before = Pt(1)
        pb.paragraph_format.space_after  = Pt(1)
        pb.paragraph_format.left_indent  = Inches(0.25)
        r = pb.add_run(bullet)
        _font(r, 10, color=DARK_GREY)


def _article(doc, title: str, published: str, summary: str,
             why: str, signals: str, source_url: str):
    """Render one GPT-4o article entry."""
    p_title = _para(doc, space_before=10, space_after=2)
    r = p_title.add_run(title)
    _font(r, 11, bold=True, color=DARK_GREY)

    if published:
        p = _para(doc, space_before=0, space_after=2)
        rl = p.add_run("Published:  ")
        _font(rl, 10, bold=True, color=MID_GREY)
        rv = p.add_run(published)
        _font(rv, 10, italic=True, color=MID_GREY)

    if summary:
        # Strip any raw HTML tags that may come through from the deterministic fallback
        clean_summary = re.sub(r"<[^>]+>", "", summary).strip()
        clean_summary = re.sub(r"&[a-zA-Z]+;|&#\d+;", " ", clean_summary).strip()
        clean_summary = re.sub(r"\s{2,}", " ", clean_summary)
        if clean_summary:
            p = _para(doc, space_before=2, space_after=3)
            rl = p.add_run("Summary:  ")
            _font(rl, 10.5, bold=True, color=DARK_GREY)
            rb = p.add_run(clean_summary)
            _font(rb, 10.5, color=DARK_GREY)

    if why:
        p = _para(doc, space_before=2, space_after=3)
        rl = p.add_run("Why it matters for ARK:  ")
        _font(rl, 10.5, bold=True, color=ACCENT_AMBER)
        rb = p.add_run(why)
        _font(rb, 10.5, color=DARK_GREY)

    if signals:
        p = _para(doc, space_before=2, space_after=3)
        rl = p.add_run("Signals to watch:  ")
        _font(rl, 10.5, bold=True, color=MID_GREY)
        rb = p.add_run(signals)
        _font(rb, 10.5, color=DARK_GREY)

    if source_url:
        p = _para(doc, space_before=2, space_after=6)
        rl = p.add_run("Source:  ")
        _font(rl, 10, bold=True, color=MID_GREY)
        _hyperlink(p, source_url, source_url)


# ── Baseline loader ───────────────────────────────────────────────────────────

def _load_baseline(path: str) -> Dict:
    p = Path(path)
    if not p.exists():
        return {}
    with open(path, encoding="utf-8") as f:
        return yaml.safe_load(f) or {}


def _active_entries(baseline: Dict, section_key: str) -> List[Dict]:
    section = baseline.get("sections", {}).get(section_key, {})
    return [
        e for e in section.get("entries", [])
        if e.get("status") in ("active",)
    ]


# ── Digest parser ─────────────────────────────────────────────────────────────

_PLACEHOLDER_PATTERNS = re.compile(
    r"no (specific|items|retrievable)|no .*(available|retrieved|selected|found)|"
    r"not available|could not be retrieved|_no in-range|_no changes",
    re.IGNORECASE,
)


def _parse_digest(md_text: str) -> Dict:
    """
    Parse ARK digest markdown — handles BOTH formats:
      New (v1.1+): ## SECTION: grants_funding  /  ### Updates This Month
      Legacy (pre-v1.1): ## Grants & Funding  (articles inline, no sub-sections)

    Returns:
    {
      'title': str,
      'exec_summary': [str],   # empty list if only placeholders
      'sections': {
        'grants_funding': {
          'articles': [{'title','published','summary','why','signals','source'}],
          'changes':  [str],
        },
        ...
      }
    }
    """
    # Strip BASELINE_DELTA block
    md_text = re.sub(
        r"---BASELINE_DELTA_START---.*?---BASELINE_DELTA_END---",
        "",
        md_text,
        flags=re.DOTALL,
    )

    result: Dict = {
        "title":        "",
        "exec_summary": [],
        "sections":     {k: {"articles": [], "changes": []} for k in _SECTION_ORDER},
    }

    lines    = md_text.splitlines()
    cur_sec  = None   # section key string
    cur_sub  = None   # "updates" | "changes" | "legacy" (old flat format)
    cur_art  = None   # dict being built

    def _flush_art():
        nonlocal cur_art
        if cur_sec and cur_art and cur_art.get("title"):
            # Only add if it has real content (not a placeholder entry)
            if not _PLACEHOLDER_PATTERNS.search(cur_art.get("title", "")):
                result["sections"][cur_sec]["articles"].append(dict(cur_art))
        cur_art = None

    in_exec = False
    i = 0
    while i < len(lines):
        raw = lines[i]
        s   = raw.strip()

        # Document title (# heading)
        if s.startswith("# ") and not s.startswith("## "):
            result["title"] = s[2:].strip()
            i += 1; continue

        # Executive summary
        if s == "**Executive Summary**":
            in_exec = True
            i += 1; continue
        if in_exec and s.startswith("- "):
            bullet = s[2:].strip()
            # Skip placeholder bullets
            if not _PLACEHOLDER_PATTERNS.search(bullet):
                result["exec_summary"].append(bullet)
            i += 1; continue

        # NEW FORMAT section marker: ## SECTION: grants_funding
        sec_m = re.match(r"^##\s+SECTION:\s+(\w+)$", s)
        if sec_m:
            _flush_art()
            in_exec = False
            sec_key = sec_m.group(1)
            cur_sec = sec_key if sec_key in _SECTION_ORDER else None
            cur_sub = "legacy"  # default until we see ### subsection
            i += 1; continue

        # LEGACY FORMAT section header: ## Grants & Funding
        if s.startswith("## ") and not sec_m:
            heading_lc = s[3:].strip().lower()
            _flush_art()
            in_exec = False
            matched_key = _LEGACY_SECTION_MAP.get(heading_lc)
            if matched_key:
                cur_sec = matched_key
                cur_sub = "legacy"
            else:
                # Unknown ## section (e.g. "## Cleantech & Start-up Ecosystem")
                # Stop collecting into any ARK section to prevent content bleed
                cur_sec = None
                cur_sub = None
            i += 1; continue

        # Sub-section markers (new format only)
        if s == "### Updates This Month":
            _flush_art()
            cur_sub = "updates"
            i += 1; continue
        if s == "### Changes Since Last Issue":
            _flush_art()
            cur_sub = "changes"
            i += 1; continue

        # Any other ### heading — stop collecting
        if s.startswith("### "):
            _flush_art()
            cur_sub = None
            i += 1; continue

        # Horizontal rule / separator
        if s == "---":
            in_exec = False
            i += 1; continue

        # Article title: bold line when in an article-collecting sub-section
        collecting_articles = cur_sec and cur_sub in ("updates", "legacy")
        if (collecting_articles
                and s.startswith("**") and s.endswith("**")
                and len(s) > 4
                and not s.startswith("**Executive")):
            _flush_art()
            cur_art = {
                "title":     s[2:-2].strip(),
                "published": "", "summary": "",
                "why":       "", "signals": "", "source": "",
            }
            i += 1; continue

        # Article fields
        if cur_art is not None:
            for prefix, field in (
                ("Published:", "published"),
                ("Summary:", "summary"),
                ("Why it matters for ARK:", "why"),
                ("Why it matters:", "why"),
                ("Signals to watch:", "signals"),
                ("Sources:", "source"),    # legacy uses plural "Sources:"
                ("Source:", "source"),
            ):
                if s.startswith(prefix):
                    val = s[len(prefix):].strip()
                    # For source/sources, take only the first URL if multiple
                    if field == "source" and val.startswith("http"):
                        val = val.split()[0]
                    cur_art[field] = val
                    break

        # Change lines (new format changes sub-section)
        if cur_sec and cur_sub == "changes" and s and not s.startswith("#"):
            if not _PLACEHOLDER_PATTERNS.search(s):
                result["sections"][cur_sec]["changes"].append(s)

        i += 1

    _flush_art()
    return result


# ── Main builder ──────────────────────────────────────────────────────────────

def build_newsletter(md_path: Path, out_path: Path,
                     baseline_path: Optional[str] = None) -> None:
    md_text  = md_path.read_text(encoding="utf-8")
    brief    = _parse_digest(md_text)
    baseline = _load_baseline(baseline_path or CFG_BASELINE)

    # Derive period label from title or filename
    period_label = ""
    title_m = re.search(r"ARK Intelligence Brief\s+—\s+(.+)", brief.get("title", ""))
    if title_m:
        period_label = title_m.group(1).strip()
    if not period_label:
        period_label = md_path.stem.replace("monthly-digest-", "")

    doc = Document()

    # Page setup — A4
    for sec in doc.sections:
        sec.page_width   = _cm_to_emu(21.0)
        sec.page_height  = _cm_to_emu(29.7)
        sec.top_margin   = _cm_to_emu(2.0)
        sec.bottom_margin= _cm_to_emu(2.0)
        sec.left_margin  = _cm_to_emu(2.5)
        sec.right_margin = _cm_to_emu(2.5)

    # Default style
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)
    style.font.color.rgb = DARK_GREY

    # ── Letterhead ─────────────────────────────────────────────────────────────
    p_co = doc.add_paragraph()
    p_co.paragraph_format.space_before = Pt(0)
    p_co.paragraph_format.space_after  = Pt(2)
    r = p_co.add_run("ARK Capture Solutions")
    _font(r, 22, bold=True, color=GG_NAVY)

    p_sub = doc.add_paragraph()
    p_sub.paragraph_format.space_before = Pt(0)
    p_sub.paragraph_format.space_after  = Pt(2)
    r = p_sub.add_run("Monthly Intelligence Brief — Australia & APAC")
    _font(r, 11, italic=True, color=GG_TEAL)

    p_prep = doc.add_paragraph()
    p_prep.paragraph_format.space_before = Pt(0)
    p_prep.paragraph_format.space_after  = Pt(6)
    r = p_prep.add_run(
        f"Prepared by GG Advisory  |  {period_label}  |  Confidential"
    )
    _font(r, 9, color=MID_GREY)

    _hrule(doc, color_hex="1B7A6B", thickness=20)

    # ── Document title ─────────────────────────────────────────────────────────
    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_before = Pt(8)
    p_title.paragraph_format.space_after  = Pt(2)
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_text = brief.get("title") or f"ARK Intelligence Brief — {period_label}"
    r = p_title.add_run(title_text)
    _font(r, 16, bold=True, color=GG_NAVY)

    _hrule(doc, color_hex="1B7A6B", thickness=8)

    # ── Executive Summary ──────────────────────────────────────────────────────
    _subsection_heading(doc, "Executive Summary", size_pt=13, color=GG_NAVY)

    # Context paragraphs from baseline
    exec_ctx = baseline.get("executive_summary", {}).get("context", {})
    for para_text in exec_ctx.get("paragraphs", []):
        p = _para(doc, space_before=2, space_after=4)
        r = p.add_run(para_text)
        _font(r, 10, color=DARK_GREY)

    # Key takeaways from digest
    p_kd = _para(doc, space_before=6, space_after=2)
    r = p_kd.add_run("Key Developments This Month")
    _font(r, 10.5, bold=True, color=GG_TEAL)

    bullets = [b for b in brief.get("exec_summary", [])
               if not _PLACEHOLDER_PATTERNS.search(b)
               and not re.search(r"intelligence brief for ark|sections:.*grants|limited high-signal", b, re.IGNORECASE)]
    if bullets:
        for bt in bullets:
            pb = doc.add_paragraph(style="List Bullet")
            pb.paragraph_format.space_before = Pt(2)
            pb.paragraph_format.space_after  = Pt(2)
            r = pb.add_run(bt)
            _font(r, 10, color=DARK_GREY)
    else:
        _shaded_para(doc,
            "— No highlights available this period. —",
            italic=True)

    _hrule(doc, color_hex="1B7A6B", thickness=6)

    # ── Sections ───────────────────────────────────────────────────────────────
    for section_key in _SECTION_ORDER:
        display_name = _SECTION_DISPLAY.get(section_key, section_key)
        _section_heading(doc, display_name)

        sec_data = brief.get("sections", {}).get(section_key, {})
        articles = sec_data.get("articles", [])
        changes  = sec_data.get("changes", [])

        # Market Background (from baseline)
        _subsection_heading(doc, "Market Background")
        active = _active_entries(baseline, section_key)
        if active:
            for entry in active:
                _baseline_entry(doc, entry)
        else:
            _shaded_para(doc, "— Baseline entries not available. —", italic=True)

        # Updates This Month (from digest)
        _subsection_heading(doc, "Updates This Month")
        if articles:
            for art in articles:
                _article(
                    doc,
                    title=art.get("title", ""),
                    published=art.get("published", ""),
                    summary=art.get("summary", ""),
                    why=art.get("why", ""),
                    signals=art.get("signals", ""),
                    source_url=art.get("source", ""),
                )
        else:
            p = _para(doc, space_before=4, space_after=4)
            r = p.add_run("No new items identified this period.")
            _font(r, 10, italic=True, color=MID_GREY)

        # Changes Since Last Issue
        _subsection_heading(doc, "Changes Since Last Issue")
        # Determine if this is truly the inaugural issue vs. just no changes this period
        _is_inaugural = not changes or all(
            re.search(r"inaugural", c, re.IGNORECASE) for c in changes
            if c.strip()
        )
        _no_changes = not _is_inaugural and all(
            c.strip().startswith("_") and c.strip().endswith("_") for c in changes
        )
        if _is_inaugural:
            _shaded_para(doc,
                "— Inaugural issue. Changes section will be populated from the second issue onwards. —",
                italic=True)
        elif _no_changes:
            p = _para(doc, space_before=4, space_after=4)
            r = p.add_run("No changes detected versus the previous issue.")
            _font(r, 10, italic=True, color=MID_GREY)
        else:
            for ch_line in changes:
                clean = ch_line.strip().lstrip("- ").lstrip("CHANGE:").strip()
                if clean and not clean.startswith("_No changes"):
                    pb = doc.add_paragraph(style="List Bullet")
                    pb.paragraph_format.space_before = Pt(1)
                    pb.paragraph_format.space_after  = Pt(2)
                    r = pb.add_run(clean)
                    _font(r, 10, color=DARK_GREY)

        _hrule(doc, color_hex="1B7A6B", thickness=6)

    # ── Footer ─────────────────────────────────────────────────────────────────
    _para(doc, space_before=10, space_after=0)
    _hrule(doc, color_hex="AAAAAA", thickness=6)
    p_footer = _para(doc, space_before=4, space_after=2)
    r = p_footer.add_run(
        "© GG Advisory  |  gg-advisory.org  |  "
        "This brief is prepared exclusively for ARK Capture Solutions. "
        "Confidential — not for redistribution."
    )
    _font(r, 8.5, italic=True, color=LIGHT_GREY)

    p_gen = _para(doc, space_before=0, space_after=0)
    r = p_gen.add_run(f"Generated: {datetime.now().strftime('%d %B %Y')}")
    _font(r, 8.5, color=LIGHT_GREY)

    doc.save(out_path)
    print(f"[ark-newsletter] Saved: {out_path}")


# ── CLI ────────────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    md_in  = Path(sys.argv[1]) if len(sys.argv) > 1 else Path("out/ark/monthly-digest-2026-03.md")
    doc_out = Path(sys.argv[2]) if len(sys.argv) > 2 else md_in.with_name(
        md_in.stem.replace("monthly-digest", "ark-intelligence-brief") + ".docx"
    )
    build_newsletter(md_in, doc_out)
