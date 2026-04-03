"""
Build a formatted Word document from the monthly digest markdown file.
Designed for direct copy-paste into the GG Advisory website blog.

Formatting choices:
- Title: Large, bold, GG Advisory teal (#1B7A6B)
- "Top Lines" intro: bold label + bullet list
- Section headings (## Energy Transition etc): bold, underlined, teal accent bar
- Each article block: headline bold, field labels bold, body text normal
- Hyperlinks on Source lines
- Generous spacing between articles for readability
"""

import re
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── Brand colours ────────────────────────────────────────────────────────────
TEAL       = RGBColor(0x1B, 0x7A, 0x6B)   # GG Advisory primary green/teal
DARK_GREY  = RGBColor(0x2C, 0x2C, 0x2C)   # near-black body text
MID_GREY   = RGBColor(0x55, 0x55, 0x55)   # secondary labels
LIGHT_TEAL = RGBColor(0xD6, 0xED, 0xE9)   # subtle section background hint (unused in text)
LINK_BLUE  = RGBColor(0x00, 0x5E, 0xA2)   # hyperlink colour


# ── Helpers ──────────────────────────────────────────────────────────────────

def _set_font(run, size_pt, bold=False, italic=False, color=None, underline=False):
    run.font.name = "Calibri"
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    run.font.underline = underline
    if color:
        run.font.color.rgb = color


def _add_paragraph(doc, space_before=0, space_after=6, left_indent=0):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after  = Pt(space_after)
    if left_indent:
        pf.left_indent = Inches(left_indent)
    return p


def _add_horizontal_rule(doc, color_hex="1B7A6B", thickness=12):
    """Add a coloured bottom-border on an empty paragraph to simulate a rule."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    str(thickness))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def _add_hyperlink(paragraph, url, display_text):
    """Insert a clickable hyperlink run into an existing paragraph."""
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    hyperlink.set(qn("w:history"), "1")
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "Hyperlink")
    rPr.append(rStyle)
    # override colour explicitly
    color_el = OxmlElement("w:color")
    color_el.set(qn("w:val"), "005EA2")
    rPr.append(color_el)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "22")   # 11pt
    szCs = OxmlElement("w:szCs")
    szCs.set(qn("w:val"), "22")
    rPr.append(sz); rPr.append(szCs)
    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.text = display_text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def _add_label_body(doc, label, body, label_color=None, body_color=None):
    """Add a paragraph with a bold label prefix and normal body text."""
    p = _add_paragraph(doc, space_before=2, space_after=3)
    r_label = p.add_run(label)
    _set_font(r_label, 11, bold=True, color=label_color or DARK_GREY)
    r_body = p.add_run(body)
    _set_font(r_body, 11, color=body_color or DARK_GREY)
    return p


# ── Section builder ───────────────────────────────────────────────────────────

def add_section_heading(doc, title):
    """## Section heading: bold, underlined, teal, with rule below."""
    _add_paragraph(doc, space_before=14, space_after=0)   # breathing room above
    p = _add_paragraph(doc, space_before=0, space_after=4)
    r = p.add_run(title.upper())
    _set_font(r, 13, bold=True, color=TEAL, underline=True)
    _add_horizontal_rule(doc, color_hex="1B7A6B", thickness=8)


def add_article(doc, title, published, summary, why, signals, source_url, source_display=""):
    """Render one digest article with all 6 fields."""
    # Article title — bold, slightly larger
    p_title = _add_paragraph(doc, space_before=10, space_after=2)
    r = p_title.add_run(title)
    _set_font(r, 11.5, bold=True, color=DARK_GREY)

    # Published date
    if published:
        p = _add_paragraph(doc, space_before=0, space_after=2)
        r_lbl = p.add_run("Published:  ")
        _set_font(r_lbl, 10.5, bold=True, color=MID_GREY)
        r_val = p.add_run(published)
        _set_font(r_val, 10.5, color=MID_GREY, italic=True)

    # Summary
    if summary:
        _add_label_body(doc, "Summary:  ", summary)

    # Why it matters
    if why:
        _add_label_body(doc, "Why it matters:  ", why)

    # Signals to watch
    if signals:
        _add_label_body(doc, "Signals to watch:  ", signals)

    # Source — hyperlink
    if source_url:
        p = _add_paragraph(doc, space_before=2, space_after=6)
        r_lbl = p.add_run("Source:  ")
        _set_font(r_lbl, 10.5, bold=True, color=MID_GREY)
        _add_hyperlink(p, source_url, source_display or source_url)


# ── Parse digest markdown ────────────────────────────────────────────────────

def parse_digest(md_text):
    """
    Parse the structured markdown digest into a dict:
      {
        'title': str,
        'top_lines': [str, ...],
        'sections': [
          {'heading': str, 'articles': [
            {'title', 'published', 'summary', 'why', 'signals', 'source'}
          ]}
        ]
      }
    """
    lines = md_text.splitlines()
    result = {"title": "", "top_lines": [], "sections": []}

    current_section = None
    current_article = None
    in_top_lines = False

    def _flush_article():
        if current_section is not None and current_article:
            current_section["articles"].append(dict(current_article))

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Document title  # February 2026
        if stripped.startswith("# ") and not stripped.startswith("## "):
            result["title"] = stripped[2:].strip()
            i += 1
            continue

        # Top Lines marker
        if stripped == "**Top Lines**":
            in_top_lines = True
            i += 1
            continue

        # Top Lines bullets
        if in_top_lines and stripped.startswith("- "):
            result["top_lines"].append(stripped[2:].strip())
            i += 1
            continue

        # Section divider or section heading
        if stripped == "---":
            in_top_lines = False
            i += 1
            continue

        if stripped.startswith("## "):
            in_top_lines = False
            _flush_article()
            current_article = None
            current_section = {"heading": stripped[3:].strip(), "articles": []}
            result["sections"].append(current_section)
            i += 1
            continue

        # Article title  **...**
        if stripped.startswith("**") and stripped.endswith("**") and current_section is not None:
            _flush_article()
            current_article = {
                "title": stripped[2:-2].strip(),
                "published": "", "summary": "", "why": "", "signals": "",
                "source": "", "source_text": "",
            }
            i += 1
            continue

        # Field lines inside an article
        if current_article is not None:
            if stripped.startswith("Published:"):
                current_article["published"] = stripped[len("Published:"):].strip()
            elif stripped.startswith("Summary:"):
                current_article["summary"] = stripped[len("Summary:"):].strip()
            elif stripped.startswith("Why it matters:"):
                current_article["why"] = stripped[len("Why it matters:"):].strip()
            elif stripped.startswith("Signals to watch:"):
                current_article["signals"] = stripped[len("Signals to watch:"):].strip()
            elif stripped.startswith("Source:"):
                raw = stripped[len("Source:"):].strip()
                # Extract URL and display text from markdown link syntax [text](url)
                m = re.match(r'\[([^\]]*)\]\(([^)]+)\)', raw)
                if m:
                    current_article["source_text"] = m.group(1)
                    current_article["source"] = m.group(2)
                else:
                    current_article["source_text"] = raw
                    current_article["source"] = raw

        i += 1

    _flush_article()
    return result


# ── Main ─────────────────────────────────────────────────────────────────────

def build_docx(md_path: Path, out_path: Path):
    md_text = md_path.read_text(encoding="utf-8")
    digest  = parse_digest(md_text)

    doc = Document()

    # Page margins: 2.5 cm each side (comfortable for web-paste preview)
    for section in doc.sections:
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin   = Inches(1.0)
        section.right_margin  = Inches(1.0)

    # Default paragraph font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.font.color.rgb = DARK_GREY

    # ── Document title ────────────────────────────────────────────────────────
    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after  = Pt(4)
    r = p_title.add_run("GG Advisory  |  Monthly Signals Digest")
    _set_font(r, 10, color=MID_GREY)

    p_month = doc.add_paragraph()
    p_month.paragraph_format.space_before = Pt(0)
    p_month.paragraph_format.space_after  = Pt(6)
    r = p_month.add_run(digest["title"])
    _set_font(r, 22, bold=True, color=TEAL)

    _add_horizontal_rule(doc, color_hex="1B7A6B", thickness=18)

    # ── Top Lines ─────────────────────────────────────────────────────────────
    p_tl_hdr = _add_paragraph(doc, space_before=10, space_after=4)
    r = p_tl_hdr.add_run("Top Lines")
    _set_font(r, 12, bold=True, underline=True, color=DARK_GREY)

    for bullet_text in digest["top_lines"]:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        r = p.add_run(bullet_text)
        _set_font(r, 11, color=DARK_GREY)

    # ── Sections ──────────────────────────────────────────────────────────────
    for section_data in digest["sections"]:
        add_section_heading(doc, section_data["heading"])
        for art in section_data["articles"]:
            add_article(
                doc,
                title          = art["title"],
                published      = art["published"],
                summary        = art["summary"],
                why            = art["why"],
                signals        = art["signals"],
                source_url     = art["source"],
                source_display = art.get("source_text", ""),
            )

    # ── Footer note ───────────────────────────────────────────────────────────
    _add_paragraph(doc, space_before=16, space_after=0)
    _add_horizontal_rule(doc, color_hex="AAAAAA", thickness=6)
    p_footer = _add_paragraph(doc, space_before=4, space_after=0)
    r = p_footer.add_run("© GG Advisory  |  gg-advisory.org  |  Monthly digest — for strategic informational purposes only.")
    _set_font(r, 9, italic=True, color=MID_GREY)

    doc.save(out_path)
    print(f"Saved: {out_path}")


if __name__ == "__main__":
    md_in  = Path(sys.argv[1]) if len(sys.argv) > 1 else Path("out/monthly-digest-2026-02.md")
    doc_out = md_in.with_suffix(".docx")
    build_docx(md_in, doc_out)
