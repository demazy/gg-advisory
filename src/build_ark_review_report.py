# -*- coding: utf-8 -*-
"""
GG Advisory — ARK Internal Review Report Builder

Builds a Word document for GG Advisory's internal use ONLY.
Never sent to ARK Capture Solutions.

Sections:
  1. Run Summary         — high-level stats for this month's run
  2. Auto-Applied Changes — baseline updates applied at >= 0.85 confidence
  3. Draft Entries       — new entries / updates at 0.60–0.85 (need review)
  4. Flagged Items       — low-confidence (<0.60) or contradiction flags
  5. Stable Entry Flags  — contradiction flags on stability:stable entries
  6. Tier-1 Fetch Alerts — source URLs that couldn't be fetched
  7. Staleness Alerts    — entries not verified within staleness_alert_days
  8. Clean Checks        — entries verified and unchanged

Env vars:
    CFG_BASELINE=config/ark-baseline.yaml
    DELTA_LOG_FILE=state/ark-delta-log-{YM}.json
    TIER1_RESULTS_FILE=state/ark-tier1-verify-{YM}.json
    OUT_DIR=out/ark
    YM=2026-03
"""
from __future__ import annotations

import json
import os
import sys
from datetime import date, datetime, timezone
from pathlib import Path
from typing import Any, Dict, List, Optional

import yaml
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

CFG_BASELINE    = os.getenv("CFG_BASELINE",     "config/ark-baseline.yaml")
DELTA_LOG_FILE  = os.getenv("DELTA_LOG_FILE",   "")
TIER1_RESULTS_FILE = os.getenv("TIER1_RESULTS_FILE", "")
OUT_DIR         = os.getenv("OUT_DIR",          "out/ark")
YM              = os.getenv("YM", os.getenv("START_YM", "")).strip()

# ── Palette (matches newsletter but with "internal" amber accent) ──────────
GG_TEAL    = RGBColor(0x1B, 0x7A, 0x6B)
GG_NAVY    = RGBColor(0x0D, 0x2D, 0x4A)
DARK_GREY  = RGBColor(0x2C, 0x2C, 0x2C)
MID_GREY   = RGBColor(0x55, 0x55, 0x55)
LIGHT_GREY = RGBColor(0x88, 0x88, 0x88)
AMBER      = RGBColor(0xE0, 0x8A, 0x00)
RED        = RGBColor(0xC0, 0x39, 0x2B)
GREEN      = RGBColor(0x27, 0xAE, 0x60)


# ── Low-level helpers ──────────────────────────────────────────────────────

def _cm_to_emu(cm: float) -> int:
    return int(cm * 914400 / 2.54)


def _font(run, size_pt: float, bold=False, italic=False,
          color: Optional[RGBColor] = None):
    run.font.name  = "Calibri"
    run.font.size  = Pt(size_pt)
    run.font.bold  = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color


def _para(doc, space_before=0, space_after=4, left_indent=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if left_indent:
        p.paragraph_format.left_indent = Inches(left_indent)
    return p


def _hrule(doc, color_hex="1B7A6B", thickness=10):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    str(thickness))
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), color_hex)
    pBdr.append(bot)
    pPr.append(pBdr)


def _section_heading(doc, title: str, color: Optional[RGBColor] = None):
    _para(doc, space_before=16, space_after=0)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(title.upper())
    _font(r, 12, bold=True, color=color or GG_NAVY)
    _hrule(doc, color_hex="0D2D4A", thickness=10)


def _status_badge(paragraph, text: str, color_hex: str):
    r = paragraph.add_run(f" [{text}] ")
    r.font.name = "Calibri"
    r.font.size = Pt(8)
    r.font.bold = True
    rPr  = r._r.get_or_add_rPr()
    col  = OxmlElement("w:color")
    col.set(qn("w:val"), color_hex)
    rPr.append(col)


def _bullet(doc, text: str, color: Optional[RGBColor] = None, left_indent=0.25):
    pb = doc.add_paragraph(style="List Bullet")
    pb.paragraph_format.space_before = Pt(1)
    pb.paragraph_format.space_after  = Pt(2)
    pb.paragraph_format.left_indent  = Inches(left_indent)
    r = pb.add_run(text)
    _font(r, 10, color=color or DARK_GREY)
    return pb


def _kv(doc, key: str, value: str, key_color: Optional[RGBColor] = None):
    p = _para(doc, space_before=1, space_after=2)
    rk = p.add_run(f"{key}  ")
    _font(rk, 10, bold=True, color=key_color or MID_GREY)
    rv = p.add_run(str(value))
    _font(rv, 10, color=DARK_GREY)
    return p


def _change_block(doc, change: Dict, badge_color: str = "1B7A6B"):
    """Render one delta-log change item."""
    action = change.get("action", "?")
    entry  = change.get("entry_id", "?") or "new"
    conf   = float(change.get("confidence", 0.0))
    desc   = change.get("description", "")[:300]
    url    = change.get("source_url", "")
    result = change.get("result", "")

    p = _para(doc, space_before=6, space_after=2)
    r = p.add_run(f"{entry} — {action}")
    _font(r, 10.5, bold=True, color=DARK_GREY)
    _status_badge(p, f"conf {conf:.2f}", badge_color)
    _status_badge(p, result or "?", "888888")

    if desc:
        _kv(doc, "Description:", desc)
    if url:
        _kv(doc, "Source:", url, key_color=MID_GREY)


# ── Data loaders ──────────────────────────────────────────────────────────────

def _load_json(path: str) -> Optional[Dict]:
    p = Path(path)
    if not p.exists():
        return None
    try:
        return json.loads(p.read_text(encoding="utf-8"))
    except Exception as e:
        print(f"[review_report] Could not load {path}: {e}")
        return None


def _load_baseline() -> Dict:
    p = Path(CFG_BASELINE)
    if not p.exists():
        return {}
    with open(CFG_BASELINE, encoding="utf-8") as f:
        return yaml.safe_load(f) or {}


def _stale_entries(baseline: Dict, threshold_days: int) -> List[Dict]:
    """Return active entries whose last_verified is older than threshold_days."""
    stale = []
    today = date.today()
    for section_key, section in baseline.get("sections", {}).items():
        for entry in section.get("entries", []):
            if entry.get("status") != "active":
                continue
            lv = entry.get("last_verified")
            if not lv:
                stale.append({"entry": entry, "section": section_key,
                               "days_since": None})
                continue
            try:
                lv_date = date.fromisoformat(str(lv)[:10])
                days    = (today - lv_date).days
                if days >= threshold_days:
                    stale.append({"entry": entry, "section": section_key,
                                  "days_since": days})
            except Exception:
                stale.append({"entry": entry, "section": section_key,
                               "days_since": None})
    return stale


def _contradiction_flags(baseline: Dict) -> List[Dict]:
    """Collect contradiction flags from all entries."""
    flags = []
    for section_key, section in baseline.get("sections", {}).items():
        for entry in section.get("entries", []):
            for flag in entry.get("contradiction_flags", []):
                flags.append({
                    "entry":   entry,
                    "section": section_key,
                    "flag":    flag,
                })
    return flags


def _draft_entries(baseline: Dict) -> List[Dict]:
    """Return all entries with status: draft or pending_updates."""
    drafts = []
    for section_key, section in baseline.get("sections", {}).items():
        for entry in section.get("entries", []):
            if entry.get("status") == "draft" or entry.get("pending_updates"):
                drafts.append({"entry": entry, "section": section_key})
    return drafts


# ── Report builder ────────────────────────────────────────────────────────────

def build_review_report(out_path: Path) -> None:
    ym       = YM or datetime.now(timezone.utc).strftime("%Y-%m")
    baseline = _load_baseline()

    # Resolve file paths
    delta_file = DELTA_LOG_FILE or f"state/ark-delta-log-{ym}.json"
    t1_file    = TIER1_RESULTS_FILE or f"state/ark-tier1-verify-{ym}.json"

    delta_log = _load_json(delta_file)
    tier1     = _load_json(t1_file)

    threshold_days = baseline.get("meta", {}).get("staleness_alert_days", 90)
    stale_list     = _stale_entries(baseline, threshold_days)
    contra_flags   = _contradiction_flags(baseline)
    draft_list     = _draft_entries(baseline)

    doc = Document()
    for sec in doc.sections:
        sec.page_width   = int(21.0 * 914400 / 2.54)
        sec.page_height  = int(29.7 * 914400 / 2.54)
        sec.top_margin   = _cm_to_emu(2.0)
        sec.bottom_margin= _cm_to_emu(2.0)
        sec.left_margin  = _cm_to_emu(2.5)
        sec.right_margin = _cm_to_emu(2.5)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)
    style.font.color.rgb = DARK_GREY

    # ── Letterhead ─────────────────────────────────────────────────────────────
    p = doc.add_paragraph()
    r = p.add_run("GG Advisory — ARK Intelligence Pipeline")
    _font(r, 18, bold=True, color=GG_NAVY)
    p.paragraph_format.space_after = Pt(2)

    p = doc.add_paragraph()
    r = p.add_run(f"Internal Review Report  |  Period: {ym}  |  CONFIDENTIAL — NOT FOR ARK")
    _font(r, 10, italic=True, color=RED)
    p.paragraph_format.space_after = Pt(6)

    _hrule(doc, color_hex="C0392B", thickness=16)

    # ── Section 1: Run Summary ─────────────────────────────────────────────────
    _section_heading(doc, "1. Run Summary")

    if delta_log:
        s = delta_log.get("summary", {})
        _kv(doc, "Period:",         delta_log.get("period", ym))
        _kv(doc, "Generated at:",   delta_log.get("generated_at", "n/a"))
        _kv(doc, "Auto-applied:",   str(s.get("auto_applied", 0)),
            key_color=GREEN)
        _kv(doc, "Draft/pending:",  str(s.get("draft_pending", 0)),
            key_color=AMBER)
        _kv(doc, "Flagged (low conf):", str(s.get("flagged", 0)),
            key_color=RED)
        _kv(doc, "Skipped:",        str(s.get("skipped", 0)))
    else:
        p = _para(doc, space_before=4)
        r = p.add_run("Delta log not found — pipeline may not have run or produced no changes.")
        _font(r, 10, italic=True, color=MID_GREY)

    if tier1:
        ok_count  = sum(1 for r in tier1.get("results", []) if r.get("fetch_status") == "ok")
        err_count = len(tier1.get("results", [])) - ok_count
        _kv(doc, "Tier-1 verified:",  f"{ok_count} OK  /  {err_count} errors")
    else:
        _kv(doc, "Tier-1 results:",   "Not available")

    _kv(doc, "Stale entries:",    str(len(stale_list)))
    _kv(doc, "Contradiction flags:", str(len(contra_flags)))
    _kv(doc, "Draft entries:",    str(len(draft_list)))

    _hrule(doc, color_hex="1B7A6B", thickness=6)

    # ── Section 2: Auto-Applied Changes ───────────────────────────────────────
    _section_heading(doc, "2. Auto-Applied Changes (≥ 0.85 confidence)")
    auto_applied = [
        c for c in (delta_log.get("changes", []) if delta_log else [])
        if c.get("gate") == "auto_apply"
        and not (c.get("result") or "").startswith("skipped")
    ]
    if auto_applied:
        for change in auto_applied:
            _change_block(doc, change, badge_color="27AE60")
    else:
        p = _para(doc, space_before=4)
        r = p.add_run("No changes auto-applied this period.")
        _font(r, 10, italic=True, color=MID_GREY)

    _hrule(doc, color_hex="1B7A6B", thickness=6)

    # ── Section 3: Draft Entries (need review) ────────────────────────────────
    _section_heading(doc, "3. Draft Entries — Pending Review")
    if draft_list:
        p0 = _para(doc, space_before=0, space_after=4)
        r = p0.add_run(
            "ACTION REQUIRED: Review each draft entry below. "
            "Promote to 'active' by changing status in ark-baseline.yaml, "
            "or delete if incorrect."
        )
        _font(r, 10, italic=True, color=AMBER)

        for item in draft_list:
            entry  = item["entry"]
            sec_k  = item["section"]
            p = _para(doc, space_before=8, space_after=2)
            r = p.add_run(f"[{sec_k}]  {entry.get('id', '?')} — {entry.get('label', '')[:80]}")
            _font(r, 10.5, bold=True, color=DARK_GREY)
            _status_badge(p, entry.get("status", "?"), "E0A800")

            # Pending updates
            for pu in entry.get("pending_updates", []):
                _kv(doc, "Pending update:",
                    f"({pu.get('action','?')}, conf {pu.get('confidence',0):.2f}) "
                    f"{pu.get('bullet','')[:150]}")
                _kv(doc, "Source:", pu.get("source_url", ""))

            # Bullets for new draft entries
            if entry.get("status") == "draft":
                for b in entry.get("bullets", []):
                    _bullet(doc, b, color=DARK_GREY)
    else:
        p = _para(doc, space_before=4)
        r = p.add_run("No draft entries this period.")
        _font(r, 10, italic=True, color=MID_GREY)

    _hrule(doc, color_hex="1B7A6B", thickness=6)

    # ── Section 4: Flagged Items (< 0.60 confidence) ──────────────────────────
    _section_heading(doc, "4. Flagged Items — Low Confidence")
    flagged = [
        c for c in (delta_log.get("changes", []) if delta_log else [])
        if c.get("gate") == "flag_only"
    ]
    if flagged:
        p0 = _para(doc, space_before=0, space_after=4)
        r = p0.add_run(
            "These changes were detected but confidence was too low to apply automatically. "
            "Review and manually update ark-baseline.yaml if warranted."
        )
        _font(r, 10, italic=True, color=AMBER)
        for change in flagged:
            _change_block(doc, change, badge_color="E67E22")
    else:
        p = _para(doc, space_before=4)
        r = p.add_run("No low-confidence flags this period.")
        _font(r, 10, italic=True, color=MID_GREY)

    _hrule(doc, color_hex="1B7A6B", thickness=6)

    # ── Section 5: Stable Entry Contradiction Flags ───────────────────────────
    _section_heading(doc, "5. Stable Entry Contradiction Flags")
    if contra_flags:
        p0 = _para(doc, space_before=0, space_after=4)
        r = p0.add_run(
            "WARNING: A source appears to contradict a stability:stable baseline entry. "
            "These entries are never auto-updated. Manually review and update if correct."
        )
        _font(r, 10, italic=True, color=RED)
        for item in contra_flags:
            entry  = item["entry"]
            sec_k  = item["section"]
            flag   = item["flag"]
            p = _para(doc, space_before=8, space_after=2)
            r = p.add_run(
                f"[{sec_k}]  {entry.get('id', '?')} — {entry.get('label', '')[:80]}"
            )
            _font(r, 10.5, bold=True, color=RED)
            _kv(doc, "Flag date:",   flag.get("date", ""))
            _kv(doc, "Description:", flag.get("description", "")[:300])
            _kv(doc, "Confidence:", f"{flag.get('confidence', 0):.2f}")
            _kv(doc, "Source:", flag.get("source_url", ""))
    else:
        p = _para(doc, space_before=4)
        r = p.add_run("No contradiction flags on stable entries.")
        _font(r, 10, italic=True, color=MID_GREY)

    _hrule(doc, color_hex="1B7A6B", thickness=6)

    # ── Section 6: Tier-1 Fetch Alerts ────────────────────────────────────────
    _section_heading(doc, "6. Tier-1 Fetch Alerts")
    if tier1:
        failed = [r for r in tier1.get("results", [])
                  if r.get("fetch_status") != "ok"]
        if failed:
            p0 = _para(doc, space_before=0, space_after=4)
            r = p0.add_run(
                "The following source URLs could not be fetched. "
                "Check if the page has moved, the grant has closed, or the site is down."
            )
            _font(r, 10, italic=True, color=AMBER)
            for fr in failed:
                p = _para(doc, space_before=6, space_after=2)
                r = p.add_run(
                    f"{fr.get('entry_id','?')} — {fr.get('label','')[:60]}"
                )
                _font(r, 10.5, bold=True, color=DARK_GREY)
                _kv(doc, "Status:", fr.get("fetch_status", "?"))
                _kv(doc, "HTTP:", str(fr.get("http_status", "n/a")))
                _kv(doc, "URL:", fr.get("url", ""))
                if fr.get("error"):
                    _kv(doc, "Error:", fr["error"][:200])
        else:
            p = _para(doc, space_before=4)
            r = p.add_run("All dynamic source URLs fetched successfully.")
            _font(r, 10, italic=True, color=GREEN)
    else:
        p = _para(doc, space_before=4)
        r = p.add_run("Tier-1 verification results not available.")
        _font(r, 10, italic=True, color=MID_GREY)

    _hrule(doc, color_hex="1B7A6B", thickness=6)

    # ── Section 7: Staleness Alerts ────────────────────────────────────────────
    _section_heading(doc, f"7. Staleness Alerts (> {threshold_days} days since last verified)")
    if stale_list:
        p0 = _para(doc, space_before=0, space_after=4)
        r = p0.add_run(
            f"These entries have not been verified in over {threshold_days} days. "
            "Manually review their source URLs and update last_verified in ark-baseline.yaml."
        )
        _font(r, 10, italic=True, color=AMBER)
        for item in sorted(stale_list, key=lambda x: x.get("days_since") or 9999, reverse=True):
            entry  = item["entry"]
            days   = item["days_since"]
            p = _para(doc, space_before=6, space_after=1)
            r = p.add_run(
                f"{entry.get('id','?')} — {entry.get('label','')[:70]}"
            )
            _font(r, 10.5, bold=True, color=DARK_GREY)
            _status_badge(p,
                f"{days} days" if days is not None else "never verified",
                "E0A800")
            _kv(doc, "Last verified:", str(entry.get("last_verified", "—")))
            _kv(doc, "Source URL:",   entry.get("source_url", "—"))
    else:
        p = _para(doc, space_before=4)
        r = p.add_run(f"All entries verified within the last {threshold_days} days.")
        _font(r, 10, italic=True, color=GREEN)

    _hrule(doc, color_hex="1B7A6B", thickness=6)

    # ── Section 8: Clean Checks ────────────────────────────────────────────────
    _section_heading(doc, "8. Clean Checks")
    all_active_ids: List[str] = []
    for section in baseline.get("sections", {}).values():
        for entry in section.get("entries", []):
            if entry.get("status") == "active":
                all_active_ids.append(entry.get("id", "?"))

    stale_ids = {item["entry"].get("id") for item in stale_list}
    contra_ids= {item["entry"].get("id") for item in contra_flags}
    draft_ids = {item["entry"].get("id") for item in draft_list}
    t1_failed_ids = {
        r.get("entry_id") for r in (tier1.get("results", []) if tier1 else [])
        if r.get("fetch_status") != "ok"
    }
    flagged_ids = {c.get("entry_id") for c in flagged} if "flagged" in dir() else set()

    clean = [eid for eid in all_active_ids
             if eid not in stale_ids
             and eid not in contra_ids
             and eid not in draft_ids
             and eid not in t1_failed_ids
             and eid not in flagged_ids]

    p0 = _para(doc, space_before=0, space_after=4)
    r = p0.add_run(
        f"{len(clean)} of {len(all_active_ids)} active entries are clean "
        f"(verified, no flags, no staleness alerts)."
    )
    _font(r, 10, color=GREEN if len(clean) == len(all_active_ids) else DARK_GREY)

    for eid in clean:
        _bullet(doc, eid, color=MID_GREY)

    # ── Footer ─────────────────────────────────────────────────────────────────
    _para(doc, space_before=14, space_after=0)
    _hrule(doc, color_hex="AAAAAA", thickness=6)
    p_footer = _para(doc, space_before=4, space_after=2)
    r = p_footer.add_run(
        "GG Advisory  |  ARK Pipeline Internal Review Report  |  "
        "CONFIDENTIAL — FOR GG ADVISORY USE ONLY — NOT FOR DISTRIBUTION TO ARK"
    )
    _font(r, 8.5, italic=True, color=RED)
    p_gen = _para(doc, space_before=0, space_after=0)
    r = p_gen.add_run(f"Generated: {datetime.now().strftime('%d %B %Y %H:%M UTC')}")
    _font(r, 8.5, color=LIGHT_GREY)

    doc.save(out_path)
    print(f"[ark-review-report] Saved: {out_path}")


# ── CLI ────────────────────────────────────────────────────────────────────────

def run() -> None:
    ym      = YM or datetime.now(timezone.utc).strftime("%Y-%m")
    out_dir = Path(OUT_DIR)
    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / f"ark-review-report-{ym}.docx"
    build_review_report(out_path)


if __name__ == "__main__":
    run()
