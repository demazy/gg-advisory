"""
Build a formatted Word document for the Grants & Accelerators Radar section.
Designed for copy-paste into B12 (or any rich-text CMS) — uses bold, italic,
bullet lists, and coloured text which survive rich-text paste.

Layout per grant:
  ── Section heading (H2 style, coloured) ────────────────────────────
  Program Name                    ← bold, coloured, Heading 3 style
  TYPE  ·  LEVEL  ·  DEADLINE     ← italic meta line

  • Funding:       Up to $500,000
  • Deadline:      30 June 2026
  • Stage:         Seed to Series A
  • Administered by: Organisation

  Summary text paragraph.

  • Why it matters:  explanation (bold label)
  • Signals to watch: text (italic)
  • Source:  hyperlink

  ────────────────────────────────────────────────   ← thin rule between cards

Groups:
  1. ⏰ CLOSING SOON        — fixed deadline within 60 days
  2. 🇦🇺 NATIONAL PROGRAMS  — rolling / tbc national
  3. 📍 STATE & TERRITORY   — by state (sub-headings per state)
"""

from __future__ import annotations

import calendar
import re
import sys
from datetime import date, timedelta
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

import yaml

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


# ── Colours ───────────────────────────────────────────────────────────────────
# These match the exact colours visible on gg-advisory.org/grants-accelerators-radar
# after manual B12 colour-picker application — they transfer correctly on paste.
ORANGE    = RGBColor(0xD0, 0x02, 0x0F)   # Closing Soon — red      #D0020F
TEAL      = RGBColor(0x68, 0x4D, 0x12)   # National programs — amber #684D12
SLATE     = RGBColor(0x7A, 0x56, 0x04)   # State & Territory — amber #7A5604
DARK_GREY = RGBColor(0x2C, 0x2C, 0x2C)
MID_GREY  = RGBColor(0x6B, 0x72, 0x80)

TYPE_LABELS = {
    "grant":           "Grant",
    "repayable_grant": "Repayable Grant",
    "accelerator":     "Accelerator",
    "incubator":       "Incubator",
    "equity":          "Equity Investment",
    "debt_equity":     "Debt + Equity",
}

LEVEL_LABELS = {
    "national": "🇦🇺 National",
    "vic": "VIC", "nsw": "NSW", "qld": "QLD",
    "sa": "SA",   "wa":  "WA",  "tas": "TAS",
    "act": "ACT", "nt":  "NT",
}

GRANTS_URGENT_DAYS = 60


# ── Low-level helpers ─────────────────────────────────────────────────────────

def _set_font(run, size_pt, bold=False, italic=False, color=None, underline=False):
    run.font.name  = "Calibri"
    run.font.size  = Pt(size_pt)
    run.font.bold  = bold
    run.font.italic = italic
    run.font.underline = underline
    if color:
        run.font.color.rgb = color


def _para(doc, space_before=0, space_after=4, left_indent=0):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after  = Pt(space_after)
    if left_indent:
        pf.left_indent = Inches(left_indent)
    return p


def _add_rule(doc, color_hex="CCCCCC", thickness=4, space_before=8, space_after=8):
    """Horizontal rule via paragraph bottom border."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    str(thickness))
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), color_hex)
    pBdr.append(bot)
    pPr.append(pBdr)
    return p


def _hyperlink(para, url: str, display: str, size_pt=10.5):
    part = para.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hl     = OxmlElement("w:hyperlink")
    hl.set(qn("r:id"),      r_id)
    hl.set(qn("w:history"), "1")
    run_el = OxmlElement("w:r")
    rPr    = OxmlElement("w:rPr")
    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "Hyperlink")
    rPr.append(rStyle)
    col = OxmlElement("w:color")
    col.set(qn("w:val"), "005EA2")
    rPr.append(col)
    sz  = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(size_pt * 2)))
    rPr.append(sz)
    szC = OxmlElement("w:szCs")
    szC.set(qn("w:val"), str(int(size_pt * 2)))
    rPr.append(szC)
    run_el.append(rPr)
    t      = OxmlElement("w:t")
    t.text = display
    run_el.append(t)
    hl.append(run_el)
    para._p.append(hl)


def _apply_heading_style(doc, para, level: int):
    style_name = f"Heading {level}"
    try:
        para.style = doc.styles[style_name]
    except Exception:
        pass


# ── Bullet list helper ────────────────────────────────────────────────────────

# text_parts: list of (text, bold, italic, color | None)
TextPart = Tuple[str, bool, bool, Optional[RGBColor]]

def _bullet(doc,
            parts: List[TextPart],
            size_pt: float = 10.5,
            space_before: int = 0,
            space_after:  int = 3):
    """Add a single Word bullet-list item from a list of styled text parts."""
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    for text, bold, italic, color in parts:
        r = p.add_run(text)
        _set_font(r, size_pt, bold=bold, italic=italic, color=color)
    return p


# ── Section heading (H2) ──────────────────────────────────────────────────────

def _section_heading(doc, title: str, color: RGBColor):
    # Leading rule — visual break above each section
    _add_rule(doc, color_hex="AAAAAA", thickness=6, space_before=20, space_after=6)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(8)
    _apply_heading_style(doc, p, 2)
    # UPPERCASE so section titles are visually distinct from program names
    # even if heading font size doesn't survive paste into B12
    r = p.add_run(title.upper())
    _set_font(r, 14, bold=True, color=color)
    return p


# ── Grant card ────────────────────────────────────────────────────────────────

def _grant_card(doc, entry: Dict, accent: RGBColor, rule_hex: str):
    """
    Render one grant as a structured block using bullets for key facts.

    Structure:
      1. Program name       → Heading 3, bold, coloured
      2. Meta line          → italic: TYPE · LEVEL · DEADLINE TYPE
      [blank]
      3. Fact bullets       → • Funding / • Deadline / • Stage / • Administered by
      [blank]
      4. Description        → body paragraph
      5. Why it matters     → bullet, bold label + body
      6. Signals to watch   → bullet, bold label + italic body
      7. Source             → bullet, bold label + hyperlink
      [rule]
    """
    def _clean(k):
        return re.sub(r"\s+", " ", str(entry.get(k) or "")).strip()

    name          = _clean("name")
    admin         = _clean("admin")
    amount        = _clean("amount")
    deadline_lbl  = _clean("deadline_label")
    target_stage  = _clean("target_stage")
    description   = _clean("description")
    why           = _clean("why_it_matters")
    signals       = _clean("signals")
    url           = _clean("url")
    prog_type     = str(entry.get("type")          or "grant").lower()
    level         = str(entry.get("level")         or "national").lower()
    deadline_type = str(entry.get("deadline_type") or "rolling").lower()

    type_label  = TYPE_LABELS.get(prog_type, prog_type.replace("_", " ").title())
    level_label = LEVEL_LABELS.get(level, level.upper())
    if deadline_type == "fixed":
        dl_label = "⏰ Fixed Deadline"
    elif deadline_type == "rolling":
        dl_label = "Rolling — always open"
    else:
        dl_label = "Date TBC"

    # ── 1. Program name — Heading 3 ───────────────────────────────────────
    # ▶ marker before the name so each program has a clear visual anchor
    # even without colour; H3 style used in case B12 preserves heading sizes
    p_name = doc.add_paragraph()
    p_name.paragraph_format.space_before = Pt(16)
    p_name.paragraph_format.space_after  = Pt(2)
    _apply_heading_style(doc, p_name, 3)
    r_arrow = p_name.add_run("▶  ")
    _set_font(r_arrow, 11, bold=True, color=accent)
    r = p_name.add_run(name or "(Untitled)")
    _set_font(r, 13, bold=True, color=accent)

    # ── 2. Meta line: TYPE · LEVEL · DEADLINE ────────────────────────────
    # Bold-italic so it reads as a label row, not body text
    p_meta = _para(doc, space_before=2, space_after=6)
    r = p_meta.add_run("  ·  ".join([type_label, level_label, dl_label]))
    _set_font(r, 10, bold=True, italic=True, color=MID_GREY)

    # ── 3. Fact bullets ───────────────────────────────────────────────────
    # Bold-italic labels so the key facts scan instantly without colour
    if amount:
        _bullet(doc, [
            ("Funding:  ",  True, True,  accent),
            (amount,        True, False, DARK_GREY),
        ])
    if deadline_lbl:
        _bullet(doc, [
            ("Deadline:  ", True, True,  accent),
            (deadline_lbl,  True, False, DARK_GREY),
        ])
    if target_stage:
        _bullet(doc, [
            ("Stage:  ",    True, True,  accent),
            (target_stage,  False, False, DARK_GREY),
        ])
    if admin:
        _bullet(doc, [
            ("Administered by:  ", True,  True,  MID_GREY),
            (admin,               False, False, MID_GREY),
        ])

    # ── 4. Description ────────────────────────────────────────────────────
    if description:
        p_sum = _para(doc, space_before=8, space_after=4)
        _set_font(p_sum.add_run(description), 10.5, color=DARK_GREY)

    # ── 5. Why it matters — bullet ────────────────────────────────────────
    if why:
        _bullet(doc, [
            ("Why it matters:  ", True, True,  accent),
            (why,                False, False, DARK_GREY),
        ], space_before=4)

    # ── 6. Signals to watch — bullet, italic ──────────────────────────────
    if signals:
        _bullet(doc, [
            ("Signals to watch:  ", True,  True,  MID_GREY),
            (signals,              False,  True,  MID_GREY),
        ])

    # ── 7. Source — bullet with hyperlink ────────────────────────────────
    if url:
        p_src = doc.add_paragraph(style="List Bullet")
        p_src.paragraph_format.space_before = Pt(0)
        p_src.paragraph_format.space_after  = Pt(3)
        rl = p_src.add_run("Source:  ")
        _set_font(rl, 10.5, bold=True, color=MID_GREY)
        _hyperlink(p_src, url, url, size_pt=10.5)

    # ── Card separator ────────────────────────────────────────────────────
    _add_rule(doc, color_hex=rule_hex, thickness=4, space_before=12, space_after=4)


# ── Data loading & classification ─────────────────────────────────────────────

def _parse_date(s: Any) -> Optional[date]:
    if not s:
        return None
    try:
        return date.fromisoformat(str(s)[:10])
    except Exception:
        return None


def load_and_classify(yaml_path: Path, ym: str):
    raw     = yaml.safe_load(yaml_path.read_text(encoding="utf-8")) or {}
    entries = raw.get("grants") or []

    y, mo = int(ym[:4]), int(ym[5:7])
    last_day = date(y + 1, 1, 1) - timedelta(days=1) if mo == 12 else date(y, mo + 1, 1) - timedelta(days=1)
    ref_date = last_day + timedelta(days=7)

    urgent, national, state = [], [], []

    for e in entries:
        if not isinstance(e, dict):
            continue
        sf = _parse_date(e.get("show_from"))
        su = _parse_date(e.get("show_until"))
        dl = _parse_date(e.get("deadline"))
        dt = str(e.get("deadline_type") or "rolling").lower()
        lv = str(e.get("level")         or "national").lower()

        if sf and ref_date < sf:
            continue
        if su and ref_date > su:
            continue
        if dt == "fixed" and dl and dl < ref_date:
            continue

        if dt == "fixed" and dl:
            days = (dl - ref_date).days
            if -7 <= days <= GRANTS_URGENT_DAYS:
                urgent.append(dict(e, _deadline_obj=dl))
                continue
        (national if lv == "national" else state).append(e)

    urgent.sort(key=lambda e: e.get("_deadline_obj") or date.max)
    national.sort(key=lambda e: str(e.get("id") or ""))
    state.sort(key=lambda e: (str(e.get("level") or ""), str(e.get("id") or "")))
    return urgent, national, state


# ── Main builder ──────────────────────────────────────────────────────────────

def _format_month_year(ym: str) -> str:
    try:
        y, m = int(ym[:4]), int(ym[5:7])
        return f"{calendar.month_name[m]} {y}"
    except Exception:
        return ym


def build_grants_docx(yaml_path: Path, out_path: Path, ym: str):
    month_year = _format_month_year(ym)
    urgent, national, state_programs = load_and_classify(yaml_path, ym)

    doc = Document()
    for sec in doc.sections:
        sec.top_margin    = Inches(1.0)
        sec.bottom_margin = Inches(1.0)
        sec.left_margin   = Inches(1.0)
        sec.right_margin  = Inches(1.0)

    style = doc.styles["Normal"]
    style.font.name      = "Calibri"
    style.font.size      = Pt(11)
    style.font.color.rgb = DARK_GREY

    # ── Page header ───────────────────────────────────────────────────────
    p_kicker = _para(doc, space_before=0, space_after=2)
    r = p_kicker.add_run("GG Advisory  |  Cleantech & Start-up Ecosystem")
    _set_font(r, 10, italic=True, color=MID_GREY)

    p_ttl = _para(doc, space_before=0, space_after=4)
    r = p_ttl.add_run(f"Grants & Accelerators Radar — {month_year}")
    _set_font(r, 22, bold=True, color=TEAL)

    p_intro = _para(doc, space_before=0, space_after=6)
    r = p_intro.add_run(
        "A curated overview of active grants, accelerators, and investment programs "
        "available to Australian climate-tech founders and investors. "
        "Entries are grouped by urgency and updated each month."
    )
    _set_font(r, 11, italic=True, color=MID_GREY)

    _add_rule(doc, color_hex="1B7A6B", thickness=12, space_before=2, space_after=14)

    # ── Group 1: Closing Soon ─────────────────────────────────────────────
    if urgent:
        _section_heading(doc, "⏰  Closing Soon — Deadlines Within 60 Days", ORANGE)
        for e in urgent:
            _grant_card(doc, e, ORANGE, "D0020F")

    # ── Group 2: National Programs ────────────────────────────────────────
    if national:
        _section_heading(doc, "🇦🇺  National Programs — Open on a Rolling Basis", TEAL)
        for e in national:
            _grant_card(doc, e, TEAL, "684D12")

    # ── Group 3: State & Territory ────────────────────────────────────────
    if state_programs:
        _section_heading(doc, "📍  State & Territory Programs", SLATE)

        current_state = None
        for e in state_programs:
            lv = str(e.get("level") or "").lower()
            if lv != current_state:
                current_state = lv
                lbl = LEVEL_LABELS.get(lv, lv.upper())
                p_st = _para(doc, space_before=14, space_after=4)
                r    = p_st.add_run(f"── {lbl.upper()} ──")
                _set_font(r, 11, bold=True, italic=True, color=SLATE)
            _grant_card(doc, e, SLATE, "7A5604")

    # ── Footer ────────────────────────────────────────────────────────────
    _add_rule(doc, color_hex="AAAAAA", thickness=4, space_before=16, space_after=4)
    p_f = _para(doc, space_before=0, space_after=0)
    r   = p_f.add_run(
        f"© GG Advisory  |  gg-advisory.org  |  {month_year} edition  |  "
        "Information is indicative only — verify deadlines and eligibility directly with the administrator."
    )
    _set_font(r, 9, italic=True, color=MID_GREY)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    print(f"Saved: {out_path}")


# ── CLI ───────────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    yaml_in = Path(sys.argv[1]) if len(sys.argv) > 1 else Path("config/grants.yaml")
    ym_arg  = sys.argv[2]       if len(sys.argv) > 2 else "2026-02"
    out     = Path(f"out/grants-radar-{ym_arg}.docx")

    sys.path.insert(0, str(Path(__file__).parent))
    build_grants_docx(yaml_in, out, ym_arg)
